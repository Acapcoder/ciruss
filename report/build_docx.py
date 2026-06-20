"""Build the CIRRUS project report as a Word .docx document.

Formatting: Times New Roman, black text only, justified body, cover page,
table of contents field, chapters, third person contributions, no em dashes,
no footer page numbers, embedded UI screenshots.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "assets")
OUT = os.path.join(HERE, "CIRRUS_Report.docx")
BLACK = RGBColor(0, 0, 0)

doc = Document()

# ---- Base style: Times New Roman, 12pt, black ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK

# Heading styles -> Times New Roman, black (override default blue)
for name, size in [("Title", 28), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    try:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.color.rgb = BLACK
        st.font.size = Pt(size)
        st.font.bold = True
    except KeyError:
        pass

# Margins ~2cm
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

IMG_W = Inches(6.0)


def para(text, justify=True, italic=False, center=False, size=None, bold=False, space_after=8):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.color.rgb = BLACK
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = BLACK
    return h


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it)
        r.font.name = "Times New Roman"
        r.font.color.rgb = BLACK


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Times New Roman"
            run.font.color.rgb = BLACK
            run.font.size = Pt(10.5)
            if i == 0:
                run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(ASSETS, filename), width=IMG_W)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose Update Field to generate the table of contents."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(placeholder); r.append(fld_end)


# ============================ COVER PAGE ============================
for _ in range(3):
    doc.add_paragraph()
para("CIRRUS", center=True, justify=False, size=46, bold=True, space_after=4)
para("Smart Multi-Cloud Storage Pooling and Compression Platform", center=True, justify=False, size=15, space_after=20)
para("Project Report", center=True, justify=False, size=20, bold=True, space_after=24)
para("Submitted in partial fulfillment of the requirements of the course", center=True, justify=False, size=12, space_after=2)
para("[Course Name and Code]", center=True, justify=False, size=12, bold=True, space_after=2)
para("[Department / Institution Name]", center=True, justify=False, size=12, space_after=24)
para("Submitted by", center=True, justify=False, size=12, bold=True, space_after=2)
for m in ["Member 1: [Full Name and Roll Number]", "Member 2: [Full Name and Roll Number]",
          "Member 3: [Full Name and Roll Number]", "Member 4: [Full Name and Roll Number]"]:
    para(m, center=True, justify=False, size=12, space_after=2)
para("Session 2025 to 2026", center=True, justify=False, size=12, space_after=2)
doc.add_page_break()

# ============================ TABLE OF CONTENTS ============================
heading("Table of Contents", 1)
add_toc()
doc.add_page_break()

# ============================ 1. INTRODUCTION ============================
heading("1. Introduction", 1)
heading("1.1 Overview", 2)
para("CIRRUS is a web based platform that allows a user to connect several free cloud storage accounts, such as Google Drive and Dropbox, and use them together as a single unified storage space. Instead of treating each cloud account separately, CIRRUS pools the free quota of every connected account into one combined capacity. Every file that a user uploads is first compressed and then automatically routed to the connected cloud that currently has the largest amount of free space.")
heading("1.2 Problem Statement", 2)
para("Individual cloud providers offer only a limited amount of free storage, for example fifteen gigabytes on Google Drive and two gigabytes on Dropbox. Users frequently hold several such accounts, yet the free space remains fragmented and difficult to use as a whole. There is no convenient consumer tool that presents these scattered free allocations as one large and easy to manage storage box.")
heading("1.3 Objectives", 2)
bullets([
    "To let users connect multiple free cloud accounts through a simple and secure sign in process.",
    "To present the combined free capacity of all connected accounts as one pooled storage box.",
    "To compress files automatically and route each file to the cloud with the most available space.",
    "To organize stored files using a unified virtual folder structure independent of the physical cloud location.",
    "To protect user accounts and cloud credentials using strong cryptographic measures.",
])
heading("1.4 Scope", 2)
para("The platform supports user registration and authentication, connection of Google Drive and Dropbox accounts through the official authorization flow, compression of files, automatic storage routing, virtual folder navigation, bulk uploads, and a dashboard that reports the pooled capacity and storage savings. The system stores only metadata and encrypted credentials, while the actual file contents reside on the user owned cloud accounts.")

# ============================ 2. ARCHITECTURE ============================
heading("2. System Architecture", 1)
heading("2.1 Architectural Overview", 2)
para("CIRRUS follows a three tier architecture. The presentation tier is a single page web application built with React. The application tier is a backend service built with the FastAPI framework. The data tier is a managed PostgreSQL database hosted on Supabase. The backend communicates with external cloud providers through their official application programming interfaces to transfer the actual file data.")
table([
    ["Tier of the System"],
    ["End User Web Browser"],
    ["React Frontend (hosted on Vercel)"],
    ["FastAPI Backend (hosted on Amazon Web Services)"],
    ["Supabase PostgreSQL Database (metadata and encrypted tokens)"],
    ["Google Drive and Dropbox (actual file storage)"],
])
para("Figure 2.1: High level three tier architecture of the CIRRUS platform, ordered from the user at the top to the storage providers at the bottom.", italic=True, center=True, justify=False, size=10.5)
heading("2.2 Technology Stack", 2)
table([
    ["Layer", "Technology", "Purpose"],
    ["Frontend", "React, Vite, React Router", "User interface, routing, and client logic"],
    ["Backend", "Python, FastAPI, Uvicorn", "API, authentication, compression, and routing engine"],
    ["Database", "PostgreSQL on Supabase", "Storage of users, accounts, folders, files, and logs"],
    ["Cloud APIs", "Google Drive API, Dropbox API", "Physical upload, download, and quota reading"],
    ["Hosting", "Amazon Web Services, Vercel", "Backend server and frontend delivery"],
])
heading("2.3 Data Flow", 2)
para("When a user uploads a file, the browser sends the file to the backend together with the authentication token. The backend compresses the file, queries the live free space of every connected account, selects the account with the most available space, and uploads the compressed file to that cloud through its official interface. The backend then records the file metadata in the database and returns a direct web link to the stored file.")

# ============================ 3. DATABASE ============================
heading("3. Database Design", 1)
heading("3.1 Overview", 2)
para("The database consists of five related tables. The design isolates user identity, connected storage accounts, virtual folders, stored file metadata, and an activity audit trail. Each record uses a universally unique identifier as its primary key. Only metadata and encrypted credentials are persisted, and no file content is stored in the database.")
heading("3.2 Tables", 2)
table([
    ["Table", "Key Columns", "Description"],
    ["users", "id, email, hashed_password, full_name, created_at", "Registered user accounts with securely hashed passwords"],
    ["storage_accounts", "id, user_id, provider, display_name, credentials_json, quota_limit, used_space", "Connected cloud accounts with encrypted refresh tokens and live quota figures"],
    ["folders", "id, user_id, name, parent_id, created_at", "Virtual folder tree using a self reference for nesting"],
    ["stored_files", "id, user_id, folder_id, account_id, original_name, compression_type, original_size, compressed_size, web_link", "Metadata for every uploaded file and its physical cloud location"],
    ["audit_logs", "id, user_id, action, details, timestamp", "Record of user actions such as sign up, login, upload, and delete"],
])
heading("3.3 Relationships", 2)
bullets([
    "A user owns many storage accounts, folders, files, and audit log entries.",
    "A storage account contains many stored files.",
    "A folder contains many files and may contain child folders through the parent identifier.",
])

# ============================ 4. FEATURES ============================
heading("4. Core Features and Implementation", 1)
heading("4.1 Storage Pooling", 2)
para("The platform aggregates the free capacity of every connected account. The free space of an account is computed as the quota limit minus the used space, and the dashboard reports the sum across all accounts as a single pooled figure.")
heading("4.2 Compression and Smart Routing", 2)
para("Each uploaded file is compressed using the gzip or zip algorithm before transfer. The routing engine then selects the connected account with the greatest free space and uploads the compressed file to that account. This approach balances usage across accounts and maximizes the effective capacity of the pool.")
heading("4.3 Virtual Folders and Bulk Upload", 2)
para("Files are organized through a unified virtual folder tree maintained in the database. A single folder may contain files that physically reside on different clouds, which preserves the concept of one combined storage box. The file manager supports selecting and uploading many files at once, and each file in a batch is routed independently.")
heading("4.4 Activity Dashboard", 2)
para("The dashboard summarizes the total number of files, the combined usage and limit, the space saved through compression, and a recent activity log. It provides the user with a clear and immediate understanding of the state of the pooled storage.")

# ============================ 5. SECURITY ============================
heading("5. Security and Privacy", 1)
heading("5.1 Authentication", 2)
para("User accounts are protected with passwords that are hashed using a salted PBKDF2 function based on the SHA256 algorithm. Sessions are maintained using signed JSON Web Tokens that expire after a fixed period.")
heading("5.2 Credential Encryption", 2)
para("Cloud authorization is performed using the official authorization code flow, which means the platform never receives or stores user passwords for the connected clouds. The resulting refresh tokens are encrypted at rest using symmetric Fernet encryption before they are written to the database.")
heading("5.3 Audit Logging", 2)
para("Significant actions, including registration, login, upload, and deletion, are recorded in an audit log table associated with the user. This provides traceability and supports accountability within the system.")

# ============================ 6. UI ============================
heading("6. User Interface", 1)
para("The interface uses a clean red and white visual theme. The following figures present the principal screens of the platform.")
figure("01_signin.png", "Figure 6.1: Sign in screen for returning users.")
figure("02_signup.png", "Figure 6.2: Registration screen with agreement to the terms of service and privacy policy.")
figure("03_dashboard.png", "Figure 6.3: Dashboard showing pooled usage, compression savings, and recent activity.")
figure("04_connections.png", "Figure 6.4: Cloud connections screen for linking Google Drive and Dropbox accounts.")
figure("05_files.png", "Figure 6.5: File manager with folder navigation, search, and bulk upload.")

# ============================ 7. DEPLOYMENT ============================
heading("7. Deployment Architecture", 1)
para("The application is deployed as a live service. The frontend is built and served through Vercel. The backend runs on an Amazon Web Services instance behind a reverse proxy with a trusted certificate, which is required for secure authorization callbacks. The database is the managed PostgreSQL service provided by Supabase. Application secrets, including the database connection string and signing keys, are stored as protected environment variables and are never committed to the source repository.")

# ============================ 8. CONTRIBUTIONS ============================
heading("8. Team Contributions", 1)
para("The project was completed by a team of four members. The responsibilities were divided into four distinct layers so that the workload was balanced and the areas of ownership did not overlap.")
heading("8.1 Member 1, Backend Core and Database", 2)
para("The first member was responsible for the backend service and the data layer. This member designed the relational schema of five tables, configured the PostgreSQL database on Supabase, and implemented the database migration logic. The member also developed the core application programming interface for files, folders, and dashboard statistics, and built the compression pipeline together with the smart routing engine that selects the emptiest cloud for each upload.")
heading("8.2 Member 2, Authentication and Security", 2)
para("The second member was responsible for all matters of authentication and security. This member implemented user registration and login, the salted password hashing scheme, and the signed token based session system. The member also implemented the symmetric encryption of cloud credentials at rest and developed the audit logging mechanism that records significant user actions.")
heading("8.3 Member 3, Cloud Integration", 2)
para("The third member was responsible for the integration with external cloud providers. This member implemented the official authorization code flow for Google Drive and Dropbox, including the secure handling and refreshing of access tokens. The member also developed the storage client layer that performs upload, download, deletion, live quota reading, and the generation of direct web links to stored files.")
heading("8.4 Member 4, Frontend and Deployment", 2)
para("The fourth member was responsible for the web application and its delivery. This member built the React interface, including the dashboard, the cloud connections manager, and the file manager with folder navigation and bulk upload. The member also designed the red and white visual theme and the guidance tooltips, and carried out the deployment of the frontend, backend, and database to their respective hosting platforms.")

# ============================ 9. CONCLUSION ============================
heading("9. Conclusion and Future Work", 1)
heading("9.1 Conclusion", 2)
para("CIRRUS successfully demonstrates that fragmented free cloud allocations can be unified into a single and practical storage box. The platform combines secure authentication, encrypted credential handling, file compression, intelligent routing, and a clear user interface into a complete and deployed product.")
heading("9.2 Future Work", 2)
bullets([
    "Adding support for further providers such as OneDrive once organizational restrictions are resolved.",
    "Implementing the splitting of very large files across multiple clouds with automatic reassembly on download.",
    "Introducing file sharing between users and richer search and preview capabilities.",
])

doc.save(OUT)
print("Word document written to", OUT, "size", round(os.path.getsize(OUT) / 1024), "KB")
